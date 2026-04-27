"""Hebrew document ingest pipeline (RAG-03 ingest half).

Pipeline: filename + bytes → ext+size+text-extract gates → chunk_hebrew → embed_batch
          → delete-existing-where-filename → collection.add → IngestResult.

Pitfall #8 mitigation: re-ingest of same filename DELETES prior chunks
BEFORE adding new ones. chunks_replaced count surfaced to caller.

D-09 lock: ALLOWED_EXTS = {".md", ".txt", ".pdf", ".docx"}; MAX_BYTES = 1 MiB.
D-03 lock: all sync chromadb calls wrapped via asyncio.to_thread.
D-02 lock: embedder is injected — never construct AsyncClient here.

Feature 2: PDF extraction via pypdf (BSD-3-Clause) with python-bidi BiDi fix
           for Hebrew visual-order PDFs. DOCX extraction via python-docx (MIT).
           Both libraries are lazy-imported inside their extractor functions.
"""

from __future__ import annotations

import asyncio
import hashlib
from datetime import UTC, datetime
from typing import TYPE_CHECKING

from receptra.rag.chunker import chunk_hebrew
from receptra.rag.errors import IngestRejected
from receptra.rag.schema import IngestResult

if TYPE_CHECKING:
    from chromadb.api.models.Collection import Collection

    from receptra.rag.embeddings import BgeM3Embedder

ALLOWED_EXTS: frozenset[str] = frozenset({".md", ".txt", ".pdf", ".docx"})
MAX_BYTES: int = 1_048_576  # 1 MiB (D-09 lock)


def _validate_extension(filename: str) -> None:
    lower = filename.lower()
    if not any(lower.endswith(ext) for ext in ALLOWED_EXTS):
        raise IngestRejected(
            code="unsupported_extension",
            detail=f"Only {sorted(ALLOWED_EXTS)} accepted in v1; got {filename!r}",
        )


def _validate_size(content: bytes) -> None:
    if len(content) > MAX_BYTES:
        raise IngestRejected(
            code="file_too_large",
            detail=f"{len(content)} bytes > {MAX_BYTES} byte limit",
        )


def _decode_utf8_strict(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError as e:
        raise IngestRejected(
            code="encoding_error",
            detail="File must be UTF-8 (strict). Save as UTF-8 and re-upload.",
        ) from e


def _extract_text(filename: str, content: bytes) -> str:
    """Dispatch to format-specific extractor based on file extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_pdf(content)
    if lower.endswith(".docx"):
        return _extract_docx(content)
    return _decode_utf8_strict(content)


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF bytes. Applies BiDi fix for Hebrew visual-order PDFs."""
    import io as _io

    import pypdf  # BSD-3-Clause
    from bidi.algorithm import get_display  # LGPL-2.1 — used as library only

    try:
        reader = pypdf.PdfReader(_io.BytesIO(content))
    except Exception as exc:
        raise IngestRejected(code="encoding_error", detail=f"PDF parse failed: {exc}") from exc

    pages: list[str] = []
    for page in reader.pages:
        raw = page.extract_text(extraction_mode="layout") or ""
        if raw.strip():
            pages.append(get_display(raw))
    text = "\n\n".join(pages)
    if not text.strip():
        raise IngestRejected(
            code="encoding_error",
            detail="PDF yielded no extractable text (scanned PDF?)",
        )
    return text


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX bytes, including table cells."""
    import io as _io

    import docx  # python-docx, MIT

    try:
        doc = docx.Document(_io.BytesIO(content))
    except Exception as exc:
        raise IngestRejected(code="encoding_error", detail=f"DOCX parse failed: {exc}") from exc

    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    if not parts:
        raise IngestRejected(
            code="encoding_error",
            detail="DOCX yielded no paragraph or table text",
        )
    return "\n\n".join(parts)


async def ingest_document(
    *,
    filename: str,
    content: bytes,
    embedder: BgeM3Embedder,
    collection: Collection,
) -> IngestResult:
    """Ingest one Hebrew document (.md/.txt/.pdf/.docx) into the RAG collection.

    Args:
        filename: original filename (extension used for allowlist check;
            stored as metadata only — never used as a filesystem path).
        content: raw file bytes (max 1 MiB). .md/.txt must be UTF-8 strict;
            .pdf/.docx are parsed by pypdf / python-docx respectively.
        embedder: BgeM3Embedder instance (injected by lifespan / test).
        collection: ChromaDB Collection (injected by lifespan / test).

    Returns:
        IngestResult with chunks_added, chunks_replaced, bytes_ingested.

    Raises:
        IngestRejected: if extension, size, encoding, or content validation fails.
    """
    _validate_extension(filename)
    _validate_size(content)
    text = _extract_text(filename, content)

    doc_sha = hashlib.sha256(content).hexdigest()
    chunks = chunk_hebrew(text)
    if not chunks:
        raise IngestRejected(
            code="empty_after_chunking",
            detail="No content after Hebrew normalization + chunking",
        )

    # Delete-before-add: re-ingest of same filename replaces prior chunks cleanly.
    # RESEARCH §Pitfall 8: re-ingest WITHOUT delete doubles chunk count silently.
    # asyncio.to_thread per D-03 (sync chromadb HttpClient must not block event loop).
    existing = await asyncio.to_thread(collection.get, where={"filename": filename})
    chunks_replaced = len(existing["ids"])
    if chunks_replaced:
        await asyncio.to_thread(collection.delete, ids=existing["ids"])

    embeddings = await embedder.embed_batch([c.text for c in chunks])
    ingested_at = datetime.now(UTC).isoformat()

    await asyncio.to_thread(
        collection.add,
        ids=[f"{doc_sha[:8]}:{c.chunk_index}" for c in chunks],
        documents=[c.text for c in chunks],
        embeddings=embeddings,  # type: ignore[arg-type]
        metadatas=[
            {
                "filename": filename,
                "chunk_index": c.chunk_index,
                "char_start": c.char_start,
                "char_end": c.char_end,
                "doc_sha": doc_sha,
                "ingested_at_iso": ingested_at,
                "tenant_id": None,  # RESEARCH §Open Decision 4 v2 forward-compat
            }
            for c in chunks
        ],
    )

    return IngestResult(
        filename=filename,
        chunks_added=len(chunks),
        chunks_replaced=chunks_replaced,
        bytes_ingested=len(content),
    )


__all__ = [
    "ALLOWED_EXTS",
    "MAX_BYTES",
    "_extract_docx",
    "_extract_pdf",
    "_extract_text",
    "ingest_document",
]
