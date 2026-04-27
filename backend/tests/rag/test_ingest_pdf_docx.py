"""Tests for PDF + DOCX text extraction in ingest.py (Feature 2).

Note: plan specified tests/kb/ but module lives in receptra.rag.ingest —
      placed in tests/rag/ to match existing project structure.
"""

from __future__ import annotations

import io

import pytest


def _make_simple_pdf(text: str) -> bytes:
    """Create a minimal valid PDF with one page containing text."""
    stream_content = f"BT /F1 12 Tf 100 700 Td ({text}) Tj ET"
    content = (
        "%PDF-1.4\n"
        "1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        "2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        "3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]\n"
        "   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
        f"4 0 obj\n<< /Length {len(stream_content)} >>\n"
        f"stream\n{stream_content}\nendstream\nendobj\n"
        "5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
        "xref\n0 6\n"
        "0000000000 65535 f \n"
        "0000000009 00000 n \n"
        "0000000058 00000 n \n"
        "0000000115 00000 n \n"
        "0000000274 00000 n \n"
        "0000000400 00000 n \n"
        "trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n477\n%%EOF"
    )
    return content.encode()


def test_allowed_exts_include_pdf_docx() -> None:
    from receptra.rag.ingest import ALLOWED_EXTS

    assert ".pdf" in ALLOWED_EXTS
    assert ".docx" in ALLOWED_EXTS


def test_extract_docx_paragraphs() -> None:
    """DOCX with paragraphs extracts text correctly."""
    import docx as _docx

    buf = io.BytesIO()
    doc = _docx.Document()
    doc.add_paragraph("שעות פתיחה 9-18")
    doc.add_paragraph("יום שישי סגור")
    doc.save(buf)
    content = buf.getvalue()

    from receptra.rag.ingest import _extract_docx

    text = _extract_docx(content)
    assert "שעות פתיחה" in text
    assert "יום שישי" in text


def test_extract_docx_tables() -> None:
    """DOCX table cells are included in extraction."""
    import docx as _docx

    buf = io.BytesIO()
    doc = _docx.Document()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "שירות"
    table.rows[0].cells[1].text = "מחיר"
    doc.save(buf)
    content = buf.getvalue()

    from receptra.rag.ingest import _extract_docx

    text = _extract_docx(content)
    assert "שירות" in text
    assert "מחיר" in text


def test_extract_docx_empty_raises() -> None:
    import docx as _docx

    from receptra.rag.errors import IngestRejected
    from receptra.rag.ingest import _extract_docx

    buf = io.BytesIO()
    doc = _docx.Document()
    doc.save(buf)
    with pytest.raises(IngestRejected) as exc_info:
        _extract_docx(buf.getvalue())
    assert exc_info.value.code == "encoding_error"


def test_extract_pdf_invalid_bytes_raises() -> None:
    from receptra.rag.errors import IngestRejected
    from receptra.rag.ingest import _extract_pdf

    with pytest.raises(IngestRejected) as exc_info:
        _extract_pdf(b"not a pdf")
    assert exc_info.value.code == "encoding_error"


@pytest.mark.asyncio
async def test_ingest_document_accepts_docx() -> None:
    """ingest_document accepts DOCX bytes end-to-end (stub embedder path)."""
    from unittest.mock import AsyncMock, MagicMock

    import docx as _docx

    from receptra.rag.ingest import ingest_document

    buf = io.BytesIO()
    doc = _docx.Document()
    doc.add_paragraph("שעות פתיחה: ראשון-חמישי 9:00-18:00")
    doc.save(buf)

    embedder = AsyncMock()
    embedder.embed_batch = AsyncMock(return_value=[[0.0] * 1024])
    collection = MagicMock()
    collection.get.return_value = {"ids": [], "documents": [], "metadatas": [], "distances": []}
    collection.add = MagicMock()

    result = await ingest_document(
        filename="hours.docx",
        content=buf.getvalue(),
        embedder=embedder,
        collection=collection,
    )
    assert result.filename == "hours.docx"
    assert result.chunks_added >= 1
